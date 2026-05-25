#!/usr/bin/env python3
"""
Convert .md files to .docx using python-docx.
Usage: python3 _scripts/md2docx.py [target_dir]
  If target_dir is given, only process files in that directory.
  Otherwise, process all known .md files that have .docx counterparts.
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to docx."""
    with open(md_path, "r") as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Consolas"
    font.size = Pt(10)

    # Add heading styles
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Consolas"
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    in_code_block = False
    code_lines = []
    in_table = False

    for raw_line in lines:
        line = raw_line.rstrip()

        # Code block
        if line.startswith("```"):
            if in_code_block:
                # End code block
                text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip empty lines
        if not line:
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            continue

        # Tables
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            # Check if separator row
            if all(re.match(r"^:?---+:?$", c.strip()) for c in cells if c.strip()):
                continue
            if not in_table:
                # First row - create table
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Light Grid Accent 1"
                for i, cell_text in enumerate(cells):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)
                            run.font.bold = True
                in_table = True
            else:
                row = table.add_row()
                for i, cell_text in enumerate(cells):
                    if i < len(row.cells):
                        cell = row.cells[i]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Consolas"
                                run.font.size = Pt(9)
            continue
        else:
            in_table = False

        # Detect heading level
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Strip emoji
            text = re.sub(r"[\U0001F300-\U0001FFFF]", "", text).strip()
            doc.add_heading(text, level=level)
            continue

        # Bold markers
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        # Inline code
        line = re.sub(r"`([^`]+)`", r"\1", line)
        # Links
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)

        # Regular paragraph
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    doc.save(docx_path)
    return True


def get_md_docx_pairs():
    """Find all .md files that should have corresponding .docx files."""
    pairs = []

    # Dasar ROS2 modul (.md → .docx, same name)
    modul_dir = os.path.join(BASE_DIR, "Dasar ROS2", "modul")
    for f in sorted(os.listdir(modul_dir)):
        if f.endswith(".md") and not f.startswith("README"):
            md_path = os.path.join(modul_dir, f)
            docx_path = os.path.join(modul_dir, f.replace(".md", ".docx"))
            pairs.append((md_path, docx_path))

    # Dasar ROS2 docs
    docs_dir = os.path.join(BASE_DIR, "Dasar ROS2", "docs")
    for f in ["ERRORS.md", "disclaimer.md", "KORELASI.md"]:
        md_path = os.path.join(docs_dir, f)
        docx_path = os.path.join(docs_dir, f.replace(".md", ".docx"))
        if os.path.exists(md_path):
            pairs.append((md_path, docx_path))

    # Dasar ROS2 explore READMEs
    explore_dir = os.path.join(BASE_DIR, "Dasar ROS2", "explore")
    for root, dirs, files in os.walk(explore_dir):
        for f in files:
            if f == "README.md":
                md_path = os.path.join(root, f)
                docx_path = os.path.join(root, f.replace(".md", ".docx"))
                pairs.append((md_path, docx_path))

    # Penerapan explore READMEs + wiring
    penerapan_explore = os.path.join(
        BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "explore"
    )
    for root, dirs, files in os.walk(penerapan_explore):
        for f in files:
            if f in ("README.md", "wiring.md"):
                md_path = os.path.join(root, f)
                docx_path = os.path.join(root, f.replace(".md", ".docx"))
                pairs.append((md_path, docx_path))

    # Penerapan referensi
    penerapan_ref = os.path.join(
        BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "referensi"
    )
    for root, dirs, files in os.walk(penerapan_ref):
        for f in files:
            if f.endswith(".md"):
                md_path = os.path.join(root, f)
                docx_path = os.path.join(root, f.replace(".md", ".docx"))
                pairs.append((md_path, docx_path))

    # Penerapan docs
    penerapan_docs = os.path.join(
        BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "docs"
    )
    for f in ["ERRORS.md", "disclaimer.md"]:
        md_path = os.path.join(penerapan_docs, f)
        docx_path = os.path.join(penerapan_docs, f.replace(".md", ".docx"))
        if os.path.exists(md_path):
            pairs.append((md_path, docx_path))

    # Other READMEs and docs
    extra_patterns = [
        (os.path.join(BASE_DIR, "AGENTS.md"), os.path.join(BASE_DIR, "AGENTS.docx")),
        (
            os.path.join(BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "AGENTS.md"),
            os.path.join(BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "AGENTS.docx"),
        ),
        (
            os.path.join(BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "README.md"),
            os.path.join(BASE_DIR, "Penerapan ROS2 pada Komponen yang Umum ditemui", "README.docx"),
        ),
    ]

    # Dasar ROS2 root-level docs
    for fname in ["README.md", "Materi.txt"]:
        md_path = os.path.join(BASE_DIR, "Dasar ROS2", fname)
        docx_path = os.path.join(BASE_DIR, "Dasar ROS2", fname.replace(".md", ".docx").replace(".txt", ".docx"))
        if os.path.exists(md_path):
            pairs.append((md_path, docx_path))

    # Dasar ROS2 scripts, launch, config, tests READMEs
    for sub in ["scripts", "launch", "config", "tests", "logs"]:
        md_path = os.path.join(BASE_DIR, "Dasar ROS2", sub, "README.md")
        docx_path = os.path.join(BASE_DIR, "Dasar ROS2", sub, "README.docx")
        if os.path.exists(md_path):
            pairs.append((md_path, docx_path))

    # Dasar ROS2 src package README
    for root, dirs, files in os.walk(os.path.join(BASE_DIR, "Dasar ROS2", "src")):
        for f in files:
            if f == "README.md":
                md_path = os.path.join(root, f)
                docx_path = os.path.join(root, f.replace(".md", ".docx"))
                pairs.append((md_path, docx_path))

    return pairs


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else None

    pairs = get_md_docx_pairs()
    if target:
        pairs = [
            (md, dx)
            for md, dx in pairs
            if target.lower() in md.lower() or target.lower() in dx.lower()
        ]

    success = 0
    failed = 0
    for md_path, docx_path in pairs:
        # Remove old docx if exists
        if os.path.exists(docx_path):
            os.remove(docx_path)

        try:
            md_to_docx(md_path, docx_path)
            print(f"  ✓ {os.path.relpath(docx_path, BASE_DIR)}")
            success += 1
        except Exception as e:
            print(f"  ✗ {os.path.relpath(md_path, BASE_DIR)}: {e}")
            failed += 1

    print(f"\nDone: {success} converted, {failed} failed")


if __name__ == "__main__":
    main()
